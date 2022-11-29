from multiprocessing import context
from django.contrib import messages
from django.shortcuts import render, redirect
from django.views.generic import View
from django.views.generic import View
from django.http import HttpResponse, HttpResponseRedirect
from .forms import *
from .models import *
from passlib.hash import pbkdf2_sha256
from django.core.mail import send_mail
from django.conf import settings
from datetime import *
from abc import ABC, abstractmethod

#notification feature
from notifications.signals import notify



import uuid
import re
# added imports below for microsoft authentication
import msal
from django.contrib.auth import login, logout
from microsoft_authentication.auth.auth_utils import (
	get_sign_in_flow,
	get_token_from_code,
	get_user,
	get_django_user,
	get_logout_url,
)
from Plan_It_Teknoy import graph
import configparser
from .graph import Graph
from microsoft_authentication.auth.auth_decorators import microsoft_login_required
# for discord
import requests
import json
import pandas as pd
import websocket #pip install websocket-client
import threading

# pip install python-docx
from docx import Document
from docx.shared import Inches

import webbrowser
# pip install docx2pdf
from docx2pdf import convert
# file manipulation
import win32com.client
import os
import pythoncom


# important don't delete

config = configparser.ConfigParser()
# config.read(['config.cfg', 'config.dev.cfg'])
config['azure'] = {'clientId': '0b89f8ff-7f30-471e-9e64-408604ee8002', 
'clientSecret': 'VL98Q~MZX6QW6~yIu1x3ozto3ehJgEg0srU.JcCP',
'tenantId' : '823cde44-4433-456d-b801-bdf0ab3d41fc',
'authTenant' : '823cde44-4433-456d-b801-bdf0ab3d41fc',
'graphUserScopes' : 'User.Read Mail.Read Mail.Send' }
azure_settings = config['azure']

graph: Graph = Graph(azure_settings)



#################################### Start of user pages ###################################

# Logout current_user session
def logout(request):
	try:
		del request.session['user']
	except:
		return redirect('Plan_It_Teknoy:Home')
	return redirect('Plan_It_Teknoy:Home')

# Home and Landing Page
def home(response):
	return render(response, "Home.html", {})

# About Us Page
def about(response):
	user = graph.get_user()
	name = user['displayName']
	email = user['mail']
	user_id = user['id']
	principal_name = user['userPrincipalName']
	
	context = {
		'user_id':user_id,
		'email':email,
		'principal_name':principal_name,
		'user': user,
		'name': name,
	}
	return render(response, "About.html", context)


# Password Reset Send Mail Function
def send_forget_password_mail(email):
	code = str(uuid.uuid4())
	subject = 'Your forget password link'
	message = f'Hi, click on the link to reset your password http://127.0.0.1:8000/ForgotPassword/{code}/'
	email_from  = settings.EMAIL_HOST_USER
	recipient_list = [email]
	send_mail(subject, message, email_from, recipient_list)
	return True

# listToString for Password Checking
def listToString(s): 
	
	# initialize an empty string
	str1 = "" 
	
	# traverse in the string  
	for ele in s: 
		str1 += ele  
	
	# return string  
	return str1 


# (After Microsoft) Index View
class IndexView(View):
	def get(self, request):
		user = graph.get_user()
		name = user['displayName']
		email = user['mail']
		user_id = user['id']
		principal_name = user['userPrincipalName']
		
		context = {
			'user_id':user_id,
			'email':email,
			'principal_name':principal_name,
			'user': user,
			'name': name,
		}
		return render(request, 'Home.html', context)
		
	def post(self,request):
		if 'btnAddStudent' in request.POST:
			if request.POST.get('student_email') and request.POST.get('student_fullname') and request.POST.get('student_id'):
				
				add_user = Users()
				add_student = Students()

				id_number = request.POST.get('student_id')

				if Students.objects.filter(StudentID=id_number).count()>0 or Users.objects.filter(id_number=id_number).count()>0:
					return redirect('Plan_It_Teknoy:calendar_view') 

				else:
					add_user.id_number = request.POST.get('student_id')
					add_user.email = request.POST.get('student_email')
					add_user.save()

					fk_id_number = Users.objects.get(id_number = id_number)

					add_student.StudentID = fk_id_number
					student_full_name = request.POST.get('student_fullname')

					name_parts = student_full_name.split()

					if len(name_parts) ==2 :
						print("there is no middle name")
						first_name = name_parts[0]
						last_name = name_parts[1]
						print(f" first name - {first_name}\n  middle name -  last name - {last_name}")
					
					elif len(name_parts) == 3:
						print("there is  middle name")
						first_name = name_parts[0]
						middle_name = name_parts[1]
						last_name = name_parts[2]

					add_student.first_name = first_name
					add_student.middle_name = middle_name
					add_student.last_name = last_name
					
					print(f" first name - {first_name}\n middle name - {middle_name} \nlast name - {last_name}")

					add_student.save()

					print('Successfully Added an applicant')
					return redirect('Plan_It_Teknoy:dashboard_view') 
		
	

def microsoft_logout(request):
	logout(request)
	return HttpResponseRedirect(get_logout_url())


# discord messages announcements start
def send_json_request(ws, request):
	ws.send(json.dumps(request))

def receive_json_response(ws):
	response = ws.recv()
	if response:
		return json.loads(response)

def heartbeat(interval, ws):
	print('Heartbeat begin')
	while True:
		time.sleep(interval)
		heartbeatJSON = {
			"op": 1,
			"d": "null"
		}
		send_json_request(ws, heartbeatJSON)
		print("Heartbeat sent")

# def retrieve_messages(channelid):
# 	headers = {
# 		'authorization': ''
# 	}
# 	messages = []
# 	r = requests.get(f'https://discord.com/api/v9/channels/{channelid}/messages', headers=headers)
# 	# r = requests.get(f'https://discord.com/api/v9/channels/1037241163226296383/messages')
# 	jsonn = json.loads(r.text)
# 	for value in jsonn:
# 		messages.append(value)

# 	return messages

def image_checker(test_list):
	for i in test_list:
		if(i == 'Image'):
			print("true")
		else:
			print("false")

# class AnnouncementsView(View):
# 	def get(self, request):
# 		ws = websocket.WebSocket()
# 		ws.connect('wss://gateway.discord.gg/?v=6&encording=json')
# 		event = receive_json_response(ws)

# 		# uncomment headers for code testing

# 		headers = {
# 		'authorization': ''
# 		}
# 		messages = []
# 		message = requests.get(f'https://discord.com/api/v9/channels/1037241163226296383/messages', headers=headers)
# 		# r = requests.get(f'https://discord.com/api/v9/channels/1037241163226296383/messages')
# 		jsonn = json.loads(message.text)
# 		for value in jsonn:
# 			messages.append(value)

# 		heartbeat_interval = event['d']['heartbeat_interval'] / 1000
# 		threading._start_new_thread(heartbeat, (heartbeat_interval, ws))
# 		token = "MTAzOTA3NDY3MjQ5MjQzMzQ3OQ.GnvIRO.k668lQVuKqsInQxrDKagJAd-CuNg9EXcEunStU"
# 		payload = {
# 			'op': 2,
# 			"d": {
# 				"token": token,
# 				"properties": {
# 					"Sos": "windows",
# 					"Sbrowser": "chrome",
# 					"Sdevice": 'pc'
# 				}
# 			}
# 		}

# 		send_json_request(ws, payload)
# 		event = receive_json_response(ws)
# 		test = []
# 		number = 0
# 		for i in messages:
# 			dictionary = {}
# 			dictionary['Username'] = messages[number]['author']['username']
# 			dictionary['Message'] = messages[number]['content']
# 			time = messages[number]['timestamp']
# 			dictionary['Time'] = pd.to_datetime(time)		
# 			image = messages[number]['attachments']
# 			# list of attachments
# 			for d in image:
# 				dictionary['Image'] = d['proxy_url']
# 			test.append(dictionary)
# 			number+=1

# 		image_check = [d['Image'] for d in test if 'Image' in d]
# 		context = {
# 			'message': message,
# 			'test': test,
# 			'dictionary': dictionary,
# 			'image_check': image_check,
# 		}
# 		return render(request, 'calendarapp/announcements.html', context)

class AnnouncementsView(View):
	def get(self, request):
		return render(request, 'calendarapp/announcements.html', {})
# discord messages announcements end

class DocGenView(View):
	def get(self, request, *args, **kwargs):
		form = EventForm(request.POST or None)
		
		user = graph.get_user()
		name = user['displayName']
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)
		current_student = Students(StudentID=confirm_user_id)
		event = Event.objects.filter(StudentID=current_student.StudentID)

		# filter [Total Events, Running Events, Completed Events]
		events = Event.objects.get_all_events(StudentID=current_student.StudentID)
		running_events = Event.objects.get_running_events(StudentID=current_student.StudentID)
		completed_events = Event.objects.get_completed_events(StudentID=current_student.StudentID)
		
		# accessing all student records in the database
		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])
		
		#document table
		docs = DocumentGen.objects.all()
		context = {
					"current_user": current_user,
					"student_record" : student_record, "form":form, "event":event, "total_event": events,
					"running_events": running_events,
					"completed_events": completed_events,
					"check_teacher": check_teacher,
					"check_student": check_student,
					"docs": docs,
					'name': name,
					}

		return render(request, 'calendarapp/document-generator.html', context)

	def post(self,request):
		if request.method == 'POST':
			if 'btnAddDoc' in request.POST:
				document = Document()
				xl=win32com.client.Dispatch("Excel.Application",pythoncom.CoInitialize())


				filename = request.POST.get('docfilename')
				content = request.POST.get('doctext')
				heading = request.POST.get('docheading', 0)

				document.add_heading(heading)
				document.add_paragraph(content)
				file = document.save(filename + ".docx")
				conv = filename + ".docx"
				convert(conv)

				add_doc = DocumentGen()
				add_doc.filename = filename
				add_doc.content = content
				add_doc.save()       

				return redirect('Plan_It_Teknoy:docgen_view')
			if 'btnViewDocument' in request.POST:
				getfile = request.POST.get('docview')
				# temporary file browsing / please replace to local drive when using
				webbrowser.open_new_tab(f'A:/GitHub/M3DA1-Plant-it-Teknoy/team_m3da1_project/{getfile}')
				return redirect('Plan_It_Teknoy:docgen_view')

			if 'btnDeleteDocument' in request.POST:
				getfile = request.POST.get('docdeletepdf')
				getfile2 = request.POST.get('docdeletedocx')
				docdeletefile = request.POST.get('docdelete')
				DocumentGen.objects.filter(DocumentID = docdeletefile).delete()
				# local drive, please change appropriately
				os.remove(f'A:/GitHub/M3DA1-Plant-it-Teknoy/team_m3da1_project/{getfile}')
				os.remove(f'A:/GitHub/M3DA1-Plant-it-Teknoy/team_m3da1_project/{getfile2}')
				return redirect('Plan_It_Teknoy:docgen_view')



# Select Role Page
class SelectRoleView(View):
	def get(self, request):
		# return render(request, 'signup-role.html', {})
		return render(request, 'user-selection.html', {})

# Contact Us page
class contactView(View):
	def get(self, request): 
		if 'user' in request.session:
			current_user = request.session['user']
			students = Students.objects.filter(StudentID=current_user)     
			users = Users.objects.filter(id_number=current_user)

			context = {
				'current_user': current_user,
				'students' : students,
				'users' : users,
			}
			return render(request, 'Contact.html', context)
		else:
			return render(request, 'Contact.html', {})

	def post(self,request):
		form = ContactForm(request.POST, request.FILES)        
		if form.is_valid():
			name = request.POST.get('name')
			email = request.POST.get("email")
			message = request.POST.get('message')
			form = Contact(name = name, email = email, message = message)         
			form.save()
			return redirect('Plan_It_Teknoy:contact_view')
		else:
			print(form.errors)
			messages.info(request, 'Please complete the fields', extra_tags='try')
			return redirect('Plan_It_Teknoy:contact_view')

# Sign in page
class SignInView(View):
	def get(self, request):
		if 'user' in request.session:
			current_user = request.session['user']
			students = Students.objects.filter(StudentID=current_user)
			teachers = Teachers.objects.filter(TeacherID=current_user) 
			users = Users.objects.filter(id_number=current_user)

			context = {
				'current_user': current_user,
				'students' : students,
				'teachers' : teachers,
				'users' : users,
			}
			return render(request, 'signin.html', context)
		else:
			return render(request, 'signin.html', {})
	
	def post(self, request):
		if request.method == 'POST':
			if 'btnSignIn' in request.POST:
				id_number = request.POST.get('id_number')
				email = request.POST.get("email")
				password = request.POST.get('password')
				check_password = Users.objects.filter(id_number=id_number).values_list("password",flat=True)
				listpw = list(check_password)
				dec_password = pbkdf2_sha256.verify(password, listToString(listpw))
				check_id = Users.objects.filter(id_number=id_number)

				check_email = Users.objects.filter(email=email)
				if check_id and dec_password and check_email:
					request.session['user'] = id_number
					if Users.objects.filter(id_number=id_number).count()>0:
						return redirect('Plan_It_Teknoy:dashboard_view')
				else:
					# does not display
					messages.info(request, 'Incorrect ID Number and Email and Password!')
					return redirect('Plan_It_Teknoy:signin_view') 

			elif 'btnForgotPass' in request.POST:
				email = request.POST.get("email")

				if Users.objects.filter(email=email).count()>0:                   
					request.session['email'] = email
					send_forget_password_mail(email)
					# add text that says email sent
					return redirect('Plan_It_Teknoy:signin_view')

				# add else if email not in db
				return redirect('Plan_It_Teknoy:contact_view')

class ForgotPasswordView(View):
	def get(self, request):
		return render(request, 'forgotpass.html', {})

	def post(self, request):
		if request.method == 'POST':
			email = request.POST.get("email")

			if Users.objects.filter(email=email).count()>0:
				
				request.session['email'] = email
				send_forget_password_mail(email)
				return redirect('Plan_It_Teknoy:Home')
			return redirect('Plan_It_Teknoy:contact_view')

class ChangePasswordSentView(View):
	def get(self, request, *args, **kwargs):
		email = request.session['email']
		context = {
			'email':email,
		}
		return render(request, 'fpsent.html', context)
	
	def post(self, request, *args, **kwargs):
		if request.method == 'POST':
			id_number = request.POST.get('id_number')
			email = request.POST.get("email")
			u = Users.objects.get(email=email)
			password = request.POST.get("password")
			cpassword = request.POST.get("cpassword")
			enc_password = pbkdf2_sha256.encrypt(password, rounds=12000, salt_size=32)
			if(cpassword == password):
				u.password = enc_password
				u.save()
				return redirect('Plan_It_Teknoy:signin_view')

		else:
			return redirect('Plan_It_Teknoy:signupS_view')

# Student Sign Up
class SignUpStudentView(View):
	def get(self, request):
		if 'user' in request.session:
			current_user = request.session['user']
			students = Students.objects.filter(StudentID=current_user)
			teachers = Teachers.objects.filter(TeacherID=current_user) 
			users = Users.objects.filter(id_number=current_user)

			context = {
				'current_user': current_user,
				'students' : students,
				'teachers' : teachers,
				'users' : users,
			}
			return render(request, 'Home.html', context)
		else:
			return render(request, 'signup-student.html', {})

	def post(self, request):        
		form1 = StudentsForm(request.POST, request.FILES)        
		form2 = UsersForm(request.POST, request.FILES)       
		if form1.is_valid() or form2.is_valid():
			# Personal Information
			first_name = request.POST.get("first_name")
			middle_name = request.POST.get("middle_name")
			last_name = request.POST.get("last_name")
			gender = request.POST.get("gender")
			department = request.POST.get("department")
			program = request.POST.get("program")
			year_level = request.POST.get("year_level")
			# Contact Information
			home_address = request.POST.get("home_address")
			city_address = request.POST.get("city_address")
			contact_number = request.POST.get("contact_number")
			# Account Information
			id_number = request.POST.get("id_number")
			password = request.POST.get("password")
			confirmpassword = request.POST.get("cpassword")
			enc_password = pbkdf2_sha256.encrypt(password, rounds=12000, salt_size=32)
			email = request.POST.get("email")
			if(confirmpassword == password):
				if Students.objects.filter(StudentID=id_number).count()>0 or Users.objects.filter(id_number=id_number).count()>0:
					messages.info(request, 'ID Number already exists!')
					return redirect('Plan_It_Teknoy:signupS_view') 
				else:              
					form2 = Users(id_number = id_number, password = enc_password, email = email)
					form2.save()
					
					fk_id_number = Users.objects.get(id_number = id_number)

					form1 = Students(StudentID = fk_id_number, first_name = first_name, middle_name = middle_name, last_name = last_name, gender = gender, department = department, program = program, year_level = year_level, home_address = home_address, city_address = city_address, contact_number = contact_number)
					form1.save()

					check_user = Users.objects.filter(id_number=id_number, password=enc_password, email = email)
					if check_user:
						request.session['user'] = id_number
						return redirect('Plan_It_Teknoy:signin_view')
			else:
				messages.info(request, 'Passwords do not match!', extra_tags='signin')
				return redirect('Plan_It_Teknoy:signupS_view')
		else:
			print(form1.errors,form2.errors)
			messages.info(request, 'Account already exists! Please try another unique one.', extra_tags='try')
			return redirect('Plan_It_Teknoy:signupS_view')

# Teacher Sign Up
class SignUpTeacherView(View):
	def get(self, request):
		if 'user' in request.session:
			current_user = request.session['user']
			students = Students.objects.filter(StudentID=current_user)
			teachers = Teachers.objects.filter(TeacherID=current_user) 
			users = Users.objects.filter(id_number=current_user)

			context = {
				'current_user': current_user,
				'students' : students,
				'teachers' : teachers,
				'users' : users,
			}
			return render(request, 'Home.html', context)
		else:
			return render(request, 'signup-teacher.html', {})

	def post(self, request):        
		form1 = StudentsForm(request.POST, request.FILES)        
		form2 = UsersForm(request.POST, request.FILES)          
		if form1.is_valid() or form2.is_valid():
			# Personal Information
			first_name = request.POST.get("first_name")
			middle_name = request.POST.get("middle_name")
			last_name = request.POST.get("last_name")
			gender = request.POST.get("gender")
			department = request.POST.get("department")
			program = request.POST.get("program")
			# Contact Information
			home_address = request.POST.get("home_address")
			city_address = request.POST.get("city_address")
			contact_number = request.POST.get("contact_number")
			# Account Information
			id_number = request.POST.get("id_number")
			password = request.POST.get("password")
			confirmpassword = request.POST.get("cpassword")
			enc_password = pbkdf2_sha256.encrypt(password, rounds=12000, salt_size=32)
			email = request.POST.get("email")
			if(confirmpassword == password):
				if Teachers.objects.filter(TeacherID=id_number).count()>0 or Users.objects.filter(id_number=id_number).count()>0:
					messages.info(request, 'ID Number already exists!')
					return redirect('Plan_It_Teknoy:signupT_view') 
				else:
					form2 = Users(id_number = id_number, password = enc_password, email = email)
					form2.save()

					fk_id_number = Users.objects.get(id_number = id_number)

					form1 = Teachers(TeacherID = fk_id_number, first_name = first_name, middle_name = middle_name, last_name = last_name, gender = gender, department = department, program = program, home_address = home_address, city_address = city_address, contact_number = contact_number)
					form1.save()

					check_user = Users.objects.filter(id_number=id_number, password=enc_password, email = email)
					if check_user:
						request.session['user'] = id_number
						return redirect('Plan_It_Teknoy:signin_view')
			else:
				messages.info(request, 'Passwords do not match!', extra_tags='signin')
				return redirect('Plan_It_Teknoy:signupT_view')
		else:
			print(form1.errors,form2.errors)
			messages.info(request, 'Account already exists! Please try another unique one.', extra_tags='try')
			return redirect('Plan_It_Teknoy:signupT_view')


""" Strategy Design Pattern for Notifications Feature"""
class IStrategy(ABC):

	@abstractmethod
	def getNotifications(self):
		"""to be implement in the child class"""

class Context():

	_strategy = IStrategy

	def __init__(self, strategy:IStrategy) -> None:
		self._strategy = strategy
	
	def implementNotifications(self):
		return self._strategy.getNotifications()


class ConcreteStrategyFirstNotifications(IStrategy):

	# Notifications for event/s that is occuring/running'
	def getNotifications(self):
		user = graph.get_user()
		current_user = user['id']
		confirm_user_id = Users(id_number=current_user)

		present_time = datetime.now()
		sender = User.objects.get(id = confirm_user_id.users_temp_id)
		receiver = User.objects.get(id = confirm_user_id.users_temp_id)

		dt_string = present_time.strftime("%Y-%m-%d %H:%M")

		get_events = Event.objects.filter(StudentID=confirm_user_id,start_time__exact = dt_string)

		queryset  = Event.objects.values_list('title', flat=True).filter(StudentID=confirm_user_id,start_time__exact = dt_string)

		for eachEvent in queryset:
			eventTitle = eachEvent

		i = 0
		while i != len(get_events):
			if get_events:
				message = "has just started." 
				notify.send(sender,recipient=receiver,verb=eventTitle,description=message, timestamp = dt_string)
				i += 1

class ConcreteStrategySecondNotifications(IStrategy):

	# Notifications for event/s that just finished/done/ended
	def getNotifications(self):
		user = graph.get_user()
		current_user = user['id']
		confirm_user_id = Users(id_number=current_user)

		present_time = datetime.now()
		sender = User.objects.get(id = confirm_user_id.users_temp_id)
		receiver = User.objects.get(id = confirm_user_id.users_temp_id)

		dt_string = present_time.strftime("%Y-%m-%d %H:%M")

		get_events = Event.objects.filter(StudentID=confirm_user_id,end_time__exact = dt_string)

		queryset  = Event.objects.values_list('title', flat=True).filter(StudentID=confirm_user_id,end_time__exact = dt_string)

		for eachEvent in queryset:
			eventTitle = eachEvent

		i = 0
		while i != len(get_events):
			if get_events:
				message = "has just ended." 
				notify.send(sender,recipient=receiver,verb=eventTitle,description=message, timestamp = dt_string)
				i += 1



class CalendarViewNew(View):

	def get(self, request):

		user = graph.get_user()
		name = user['displayName']
		form = EventForm(request.POST or None)

		current_user = user['id']
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)
		confirm_user_id = Users(id_number=current_user)
		current_student = Students(StudentID=confirm_user_id)
		running_events = Event.objects.get_running_events(StudentID=current_student.StudentID)


		# Applying Notification Feature in this Class
		contextOccuringEvents = Context(ConcreteStrategyFirstNotifications())
		contextOccuringEvents.implementNotifications()


		# Applying Notification Feature in this Class
		contextEndedEvents = Context(ConcreteStrategySecondNotifications())
		contextEndedEvents.implementNotifications()


		#accessing all student records in the database
		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level, profile_pic FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])
		
		student_running_events = []

		for student_event in running_events:
			student_running_events.append(
				{
				"title":student_event.title,
				"description": student_event.description,
				"start":student_event.start_time.strftime("%Y-%m-%d %H:%M:%S"),
				"end":student_event.end_time.strftime("%Y-%m-%d %H:%M:%S"),
				}
			)

		context = {
			"student_running_events":student_running_events,  
			"running_events":running_events, 
			"form":form, "student_record":student_record, 
			"check_teacher":check_teacher, 
			"check_student": check_student, 
			'name': name,
			}

		return render(request, 'calendarapp/calendar.html', context)

	def post(self, request):
		user = graph.get_user()        
		form2 = EventForm(request.POST or None)        
		if request.POST and form2.is_valid():
			current_user = user['id']
			confirm_user_id = Users(id_number=current_user)
			current_student = Students(StudentID=confirm_user_id)
			title = form2.cleaned_data["title"]
			description = form2.cleaned_data["description"]
			start_time = form2.cleaned_data["start_time"]
			end_time = form2.cleaned_data["end_time"]
			form2 = Event(StudentID = current_student.StudentID, title = title, description = description, start_time = start_time, end_time = end_time)		
			form2.save()
			return redirect('Plan_It_Teknoy:calendar_view')
				
		else:
			print(form2.errors)
			messages.info(request, 'Form error.', extra_tags='try')
			return redirect('Plan_It_Teknoy:calendar_view')

# DashboardView
class DashboardView(View):
	def get(self, request, *args, **kwargs):
		form = EventForm(request.POST or None)
		
		user = graph.get_user()
		name = user['displayName']
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)
		current_student = Students(StudentID=confirm_user_id)
		event = Event.objects.filter(StudentID=current_student.StudentID)

		# filter [Total Events, Running Events, Completed Events]
		events = Event.objects.get_all_events(StudentID=current_student.StudentID)
		running_events = Event.objects.get_running_events(StudentID=current_student.StudentID)
		completed_events = Event.objects.get_completed_events(StudentID=current_student.StudentID)
		
		# accessing all student records in the database
		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])
		
		context = {
					"current_user": current_user,
					"student_record" : student_record, "form":form, "event":event, "total_event": events,
					"running_events": running_events,
					"completed_events": completed_events,
					"check_teacher": check_teacher,
					"check_student": check_student,
					'name': name,
					}

		return render(request, 'calendarapp/dashboard.html', context)

	def post(self, request):
		if request.method == 'POST':
			if 'btnUpdateEvent' in request.POST:
				eventID = request.POST.get("eventID")
				eventTitle = request.POST.get("eventTitle")
				eventDesc = request.POST.get("eventDesc")
				eventST = request.POST.get("eventST")
				eventET = request.POST.get("eventET")
				Event.objects.filter(EventID = eventID).update(title = eventTitle,
					description = eventDesc, start_time = eventST, end_time = eventET)

			if 'btnDeleteEvent' in request.POST:
				deleventID = request.POST.get("deleventID")
				Event.objects.filter(EventID = deleventID).delete()

		return redirect('Plan_It_Teknoy:dashboard_view')
	
class AllEventsListView(ListView):

	# """ All event list views """

	def get(self, request):

		user = graph.get_user()
		name = user['displayName']
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		current_student = Students(StudentID=confirm_user_id)
		events = Event.objects.get_all_events(StudentID=current_student.StudentID)
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)



		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])

		context = {
			"student_record" : student_record, 
			"total_event":events,
			"events":events, 
			"check_teacher":check_teacher, 
			"check_student":check_student,
			'name': name,
			}

		return render(request, 'calendarapp/events_list.html', context)
		


class RunningEventsListView(ListView):

	# """ Running events list view """

	
	def get(self, request):

		user = graph.get_user()
		name = user['displayName']
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		current_student = Students(StudentID=confirm_user_id)
		running_events = Event.objects.get_running_events(StudentID=current_student.StudentID)
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)


		# present_time = datetime.now()
		# sender = User.objects.get(id = confirm_user_id.users_temp_id)
		# receiver = User.objects.get(id = confirm_user_id.users_temp_id)
	

		# dt_string = present_time.strftime("%Y-%m-%d %H:%M")
		# get_events = Event.objects.filter(StudentID=confirm_user_id,start_time__exact = dt_string)
		# print("HAHAH: ", len(get_events))

		# i = 0
		# while i != len(get_events):
		# 	if Event.objects.filter(StudentID=confirm_user_id,start_time__exact = dt_string):
		# 		message = "One Event is already running seconds ago."
		# 		notify.send(sender,recipient=receiver,verb='Event Running',description=message, timestamp = dt_string)
		# 		i += 1
		
		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])
		
		context = {
			"student_record" : student_record, 
			"events":running_events, 
			"running_events":running_events, 
			"check_teacher":check_teacher, 
			"check_student":check_student,
			'name': name,
			}

		return render(request, 'calendarapp/events_list.html', context)

class CompletedEventsListView(ListView):

	def get(self, request):

		user = graph.get_user()
		name = user['displayName']
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		current_student = Students(StudentID=confirm_user_id)
		completed_events = Event.objects.get_completed_events(StudentID=current_student.StudentID)
		check_teacher = Teachers.objects.filter(TeacherID=current_user)
		check_student = Students.objects.filter(StudentID=current_user)

		student_record = Students.objects.raw('SELECT StudentID_id, first_name, program, last_name, year_level FROM plan_it_teknoy_students WHERE StudentID_id = %s', [current_student.StudentID])
		
		context = {
			"student_record" : student_record, 
			"events":completed_events, 
			"completed_events":completed_events, 
			"check_teacher":check_teacher, 
			"check_student":check_student,
			'name': name,
			}

		return render(request, 'calendarapp/events_list.html', context)

# class NotificationsListView(ListView):

# 	def get(self, request):
# 		user = graph.get_user()   
# 		current_user = user['id']
# 		present_time = datetime.now()
# 		confirm_user_id = Users(id_number=current_user)
	
# 		present_time = datetime.now()
# 		sender = User.objects.get(id = confirm_user_id.users_temp_id)
# 		receiver = User.objects.get(id = confirm_user_id.users_temp_id)

# 		dt_string = present_time.strftime("%Y-%m-%d %H:%M")

# 		if Event.objects.filter(StudentID=confirm_user_id,start_time__exact = dt_string):
# 			message = "One Event is already running seconds ago."
# 			notify.send(sender,recipient=receiver,verb='Event Running',description=message, timestamp = dt_string)


# 		return render(request, 'base/base.html', {})


# user_profile_settings_views
class SProfileSettings(View):

	def get(self, request):

		user = graph.get_user()
		current_user = user['id']

		confirm_user_id = Users(id_number=current_user)
		current_student = Students.objects.filter(StudentID=confirm_user_id)
		email_student = Users.objects.filter(id_number=confirm_user_id) 

		return render(request, 'account-settings.html', {"current_student":current_student, "email_student":email_student})
  
	def post(self, request):
		if request.method == 'POST':

			# profile pic update feature
			if 'btnUpdateProPic' in request.POST:
				print('UpdateProPic button clicked!')
				student_id = request.POST.get("student_id")
				profile_pic = request.FILES['profile_pic']
				saveProPic = Students.objects.get(StudentID = student_id)
				saveProPic.profile_pic = profile_pic
				saveProPic.save()
				messages.success(request, "Profile Picture Successfully Updated!!!", extra_tags='profile_pic_success')
				print('Student profile picture updated!')
				return redirect('Plan_It_Teknoy:sprofile-settings_view')   
			

			# personal details update feature
			if 'btnUpdate' in request.POST:
				print('UpdateDetails button clicked!')
				student_id = request.POST.get("student_id")
				firstname = request.POST.get("first_name")
				lastname = request.POST.get("last_name")
				cNumber = request.POST.get("contact_number")
				sGender = request.POST.get("gender")
				hAddress = request.POST.get("home_address")
				cAddress = request.POST.get("city_address")
				Students.objects.filter(StudentID = student_id).update(first_name = firstname, last_name = lastname, contact_number = cNumber, gender = sGender, home_address = hAddress, city_address = cAddress)
				messages.success(request, "Profile Details Successfully Updated!!!", extra_tags='profile_details_success')
				print('Student account updated!')

			# student academic details
			if 'updateAcademicButton' in request.POST:
				print('UpdateDetails button clicked!')
				student_id = request.POST.get("student_academic_id")
				sdepartment = request.POST.get("academic_department")
				sprogram = request.POST.get("academic_program")
				syear_level = request.POST.get("academic_year_level")
				Students.objects.filter(StudentID = student_id).update(department = sdepartment, program = sprogram, year_level = syear_level)
				messages.success(request, "Academic Details Successfully Updated!!!", extra_tags='academic_details_success')
				print('Student account academic updated!')
			

			if 'btnDeleteStudent' in request.POST:
				print('Delete button clicked!')
				student_id = request.POST.get("current_student_id")
				Students.objects.filter(StudentID=student_id).delete()
				Users.objects.filter(id_number=student_id).delete()
				Users.objects.filter(id_number=student_id).delete()
				Event.objects.filter(StudentID = student_id).delete()
				print("Student account deleted")
				return redirect('Plan_It_Teknoy:Logout')
						
			
			if 'btnSubmitPassword' in request.POST:

				student_id = request.POST.get("student_id_security")

				# get_user_id = Users.objects.get(id_number = student_id)

				user_current_pwd = request.POST.get("current_password")

				check_exact_pwd = Users.objects.filter(id_number = student_id).values_list("password", flat=True)
				pass_list = list(check_exact_pwd)
				decrypt_pass = pbkdf2_sha256.verify(user_current_pwd, listToString(pass_list))

				if decrypt_pass:

					print("Current password and newly inputted password matched!")

					user_new_pwd = request.POST.get("new_password")
					enc_user_new_pwd = pbkdf2_sha256.encrypt(user_new_pwd, rounds=12000, salt_size=32)

					user_confirm_new_pwd = request.POST.get("confirm_password")

					if user_new_pwd == user_confirm_new_pwd:
						Users.objects.filter(id_number = student_id).update(password = enc_user_new_pwd)
						print("Password newly created")
						messages.success(request, "Account Password Successfully Updated!!!", extra_tags='pass_success')
						return redirect('Plan_It_Teknoy:sprofile-settings_view')
					
					else:
						messages.error(request, 'New password and Confirm password did not matched!', extra_tags='old_new_pass_error')
						#pwede ni siya ma message para ma send sa html nya mag modal pop up

				else:
					messages.error(request, 'You did not input your correct current password!', extra_tags='current_pass_error')

			return redirect('Plan_It_Teknoy:sprofile-settings_view')


class TProfileSettings(View):
	def get(self, request):

			if 'user' in request.session:
				current_user = request.session['user']
				confirm_user_id = Users(id_number=current_user)
				current_teacher = Teachers.objects.filter(TeacherID=confirm_user_id)
				email_teacher = Users.objects.filter(id_number=confirm_user_id) 

				return render(request, 'teacher-account-settings.html', {"current_teacher":current_teacher, "email_teacher":email_teacher})
  
	def post(self, request):
		if request.method == 'POST':

			# profile pic update feature
			if 'btnUpdateProPic' in request.POST:
				print('UpdateProPic button clicked!')
				teacher_id = request.POST.get("teacher_id")
				profile_pic = request.FILES['profile_pic']
				saveProPic = Teachers.objects.get(TeacherID = teacher_id)
				saveProPic.profile_pic = profile_pic
				saveProPic.save()
				messages.success(request, "Profile Picture Successfully Updated!!!", extra_tags='profile_pic_success')
				print('Teacher profile picture updated!')
				return redirect('Plan_It_Teknoy:tprofile-settings_view')   
			

			# personal details update feature
			if 'btnUpdate' in request.POST:
				print('UpdateDetails button clicked!')
				teacher_id = request.POST.get("teacher_id")
				firstname = request.POST.get("first_name")
				lastname = request.POST.get("last_name")
				cNumber = request.POST.get("contact_number")
				sGender = request.POST.get("gender")
				hAddress = request.POST.get("home_address")
				cAddress = request.POST.get("city_address")
				Teachers.objects.filter(TeacherID = teacher_id).update(first_name = firstname, last_name = lastname, contact_number = cNumber, gender = sGender, home_address = hAddress, city_address = cAddress)
				messages.success(request, "Profile Details Successfully Updated!!!", extra_tags='profile_details_success')
				print('Teacher account updated!')

			# teacher academic details
			if 'updateAcademicButton' in request.POST:
				print('UpdateDetails button clicked!')
				teacher_id = request.POST.get("teacher_academic_id")
				sdepartment = request.POST.get("academic_department")
				sprogram = request.POST.get("academic_program")
				# syear_level = request.POST.get("academic_year_level")
				Teachers.objects.filter(TeacherID = teacher_id).update(department = sdepartment, program = sprogram)
				messages.success(request, "Academic Details Successfully Updated!!!", extra_tags='academic_details_success')
				print('Teacher account academic updated!')
			

			if 'btnDeleteTeacher' in request.POST:
				print('Delete button clicked!')
				teacher_id = request.POST.get("current_teacher_id")
				Teachers.objects.filter(TeacherID=teacher_id).delete()
				if 'user' in request.session:
					current_teacher = request.session['user']
					Users.objects.filter(id_number=current_teacher).delete()
					Event.objects.filter(TeacherID = current_teacher).delete()
				print("Teacher account deleted")
				return redirect('Plan_It_Teknoy:Logout')
						
			
			if 'btnSubmitPassword' in request.POST:

				teacher_id = request.POST.get("teacher_id_security")

				# get_user_id = Users.objects.get(id_number = student_id)

				user_current_pwd = request.POST.get("current_password")

				check_exact_pwd = Users.objects.filter(id_number = teacher_id).values_list("password", flat=True)
				pass_list = list(check_exact_pwd)
				decrypt_pass = pbkdf2_sha256.verify(user_current_pwd, listToString(pass_list))

				if decrypt_pass:

					print("Current password and newly inputted password matched!")

					user_new_pwd = request.POST.get("new_password")
					enc_user_new_pwd = pbkdf2_sha256.encrypt(user_new_pwd, rounds=12000, salt_size=32)

					user_confirm_new_pwd = request.POST.get("confirm_password")

					if user_new_pwd == user_confirm_new_pwd:
						Users.objects.filter(id_number = teacher_id).update(password = enc_user_new_pwd)
						print("Password newly created")
						messages.success(request, "Account Password Successfully Updated!!!", extra_tags='pass_success')
						return redirect('Plan_It_Teknoy:tprofile-settings_view')
					
					else:
						messages.error(request, 'New password and Confirm password did not matched!', extra_tags='old_new_pass_error')
						#pwede ni siya ma message para ma send sa html nya mag modal pop up

				else:
					messages.error(request, 'You did not input your correct current password!', extra_tags='current_pass_error')

			return redirect('Plan_It_Teknoy:tprofile-settings_view')



		

	